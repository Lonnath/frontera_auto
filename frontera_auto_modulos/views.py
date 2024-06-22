# DJANGO IMPORTS


from django.views.decorators.http import require_http_methods
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse, HttpResponse
from django.core.files.base import ContentFile
from frontera_auto_back.settings import SECRET_KEY

# UTILS IMPORTS

import base64
import json
import jwt
from users.decorators import jwt_token_required
from frontera_auto_back.settings import SECRET_KEY
import os
from docx import Document
from docx2pdf import convert
from docx.shared import Inches

# SERIALIZERS

from .serializers import *

# OBJECTS IMPORTS

from users.models import Usuario
from .models import *

# UTIL METHODS


def replace_text(run, replacements):
    for key, value in replacements.items():
        if key in run.text:
            run.text = run.text.replace(key, str(value))

# PESO LOGIC


@csrf_exempt
@require_http_methods(["POST"])
@jwt_token_required
def create_peso(request):
    auth_header = request.headers.get('Authorization')
    token = auth_header.split(' ')[1]
    payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
    data = json.loads(request.body)
    data['usuario'] = payload['id']
    serializer = PesajeSubproductoSerializer(data=data)

    if serializer.is_valid():
        pesaje_subproducto_serializer = serializer.save()
        return JsonResponse({'response': f'Peso creado correctamente para la Viscera {pesaje_subproducto_serializer.id}'}, status=200)
    else:
        return JsonResponse({'response': f'ERROR: No se pudo guardar el registro, Causa: {serializer.errors}.'}, status=400)


@csrf_exempt
@require_http_methods(["GET"])
@jwt_token_required
def get_peso(request, id=int):

    peso_data = None

    try:
        peso_data = PesajeSubproducto.objects.get(id=id)
    except PesajeSubproducto.DoesNotExist:
        return JsonResponse({'response': 'Pesaje Subproducto no encontrado'}, status=400)

    serializer = PesajeSubproductoSerializer(peso_data)
    out_data = {}
    out_data['pesaje_subproducto'] = serializer.data
    out_data['viscera'] = {
        'turno': peso_data.vicera.turno,
        'cifra': peso_data.vicera.cifra,
        'peso': peso_data.vicera.peso,
        'serie': peso_data.vicera.serie,
        'responsable_firma': peso_data.vicera.responsable_firma,
        'tipo_vicera': peso_data.vicera.tipo_vicera,

    }
    return JsonResponse({'response': 'Pesaje Subproducto consultada correctamente', 'Message': out_data}, status=200)


@csrf_exempt
@require_http_methods(["GET"])
@jwt_token_required
def get_pesos_all(request):

    peso_data = None

    try:
        peso_data = PesajeSubproducto.objects.all()
    except PesajeSubproducto.DoesNotExist:
        return JsonResponse({'response': 'Pesaje Subproducto no encontrado'}, status=400)

    serializer = PesajeSubproductoSerializer(peso_data, many=True)

    return JsonResponse({'response': 'Pesaje Subproducto consultada correctamente', 'Message': serializer.data}, status=200)


# VISCERAS LOGIC

@csrf_exempt
@require_http_methods(["POST"])
@jwt_token_required
def create_viscera(request):

    data = json.loads(request.body)
    serializer = ViceraSerializer(data=data)
    if serializer.is_valid():
        viscera_data = serializer.save()
        return JsonResponse({'response': 'Viscera creada correctamente', 'Message': {'id': viscera_data.id}}, status=200)
    else:
        return JsonResponse({'response': f'ERROR: No se pudo guardar el registro, Causa: {serializer.errors}.'}, status=400)


@csrf_exempt
@require_http_methods(["GET"])
@jwt_token_required
def get_viscera(request, id=int):

    visceras_data = None

    try:
        visceras_data = Vicera.objects.get(id=id)
    except Vicera.DoesNotExist:
        return JsonResponse({'response': 'Viscera no encontrada'}, status=400)

    serializer = ViceraSerializer(visceras_data)

    return JsonResponse({'response': 'Viscera consultada correctamente', 'Message': {'data': serializer.data}}, status=200)


@csrf_exempt
@require_http_methods(["GET"])
@jwt_token_required
def get_visceras_all(request):

    visceras_data = None
    try:
        visceras_data = Vicera.objects.all()
    except Vicera.DoesNotExist:
        return JsonResponse({'response': 'Visceras consultadas incorrectamente'}, status=400)

    serializer = ViceraSerializer(visceras_data, many=True)
    return JsonResponse({'response': 'Visceras consultadas correctamente', 'Message': {'data': serializer.data}}, status=200)

# DECOMISOS LOGIC


@csrf_exempt
@require_http_methods(["POST"])
@jwt_token_required
def create_decomiso(request):
    auth_header = request.headers.get('Authorization')
    token = auth_header.split(' ')[1]
    payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
    data = json.loads(request.body)
    data['usuario'] = payload['id']

    foto_base64 = data.pop('foto', None)
    if foto_base64:
        format, imgstr = foto_base64.split(';base64,')
        ext = format.split('/')[-1]
        photo_file = ContentFile(base64.b64decode(
            imgstr), name=f'decomiso_{data["usuario"]}.{ext}')
        data['foto'] = photo_file

    serializer = DecomisoSerializer(data=data)
    if serializer.is_valid():
        decomiso_data = serializer.save()
        return JsonResponse({'response': 'Decomiso creado correctamente', 'Message': {'id': decomiso_data.id}}, status=200)
    else:
        return JsonResponse({'response': f'ERROR: No se pudo guardar el registro, Causa: {serializer.errors}.'}, status=400)


@csrf_exempt
@require_http_methods(["GET"])
@jwt_token_required
def get_decomiso(request, id=int):

    decomiso_data = None

    try:
        decomiso_data = Decomiso.objects.get(id=id)
    except Decomiso.DoesNotExist:
        return JsonResponse({'response': 'Decomiso no encontrado'}, status=400)

    serializer = DecomisoSerializer(decomiso_data)

    return JsonResponse({'response': 'Decomiso consultado correctamente', 'Message': {'data': serializer.data}}, status=200)


@csrf_exempt
@require_http_methods(["GET"])
@jwt_token_required
def get_decomisos_all(request):

    decomisos_data = None
    try:
        decomisos_data = Decomiso.objects.all()
    except Decomiso.DoesNotExist:
        return JsonResponse({'response': 'Decomisos consultados incorrectamente'}, status=400)

    serializer = DecomisoSerializer(decomisos_data, many=True)
    return JsonResponse({'response': 'Decomisos consultados correctamente', 'Message': {'data': serializer.data}}, status=200)


@csrf_exempt
@require_http_methods(["GET"])
@jwt_token_required
def get_pdf_decomiso(request, id):
    decomiso_data = None

    try:
        decomiso_data = Decomiso.objects.get(id=id)
    except Decomiso.DoesNotExist:
        return JsonResponse({'response': 'Decomiso no encontrado'}, status=400)

    path_name = "media/Plantilla_Decomiso.docx"
    output_docx_name = f"media/Informe_Decomisos_{id}.docx"
    output_pdf_name = f"media/Informe_Decomisos_{id}.pdf"

    data = {
        'ciudad': decomiso_data.ciudad,
        'producto': decomiso_data.producto,
        'serie': decomiso_data.serie,
        'fecha': decomiso_data.fecha.strftime("%d/%m/%y"),
        'numC': decomiso_data.cifra,
        'cantidad': decomiso_data.cantidad,
        'causal': decomiso_data.causa,
        'numT': decomiso_data.turno,
        'nombre_veterinario': decomiso_data.usuario.nombre,
        'numero_targ_prod': decomiso_data.usuario.cc_documento,
        'cn': id,
        'observations': decomiso_data.observacion
    }

    # Cargar el documento de Word
    doc = Document(path_name)

    # Reemplazar las variables en los párrafos
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            replace_text(run, data)

    # Reemplazar las variables dentro de las tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        replace_text(run, data)

    image_path = os.path.join("media", decomiso_data.foto.name)

    # Añadir la imagen al documento
    for paragraph in doc.paragraphs:
        if 'foto' in paragraph.text:
            paragraph.text = paragraph.text.replace('foto', '')
            paragraph.add_run().add_picture(image_path, width=Inches(2.0))

    # Guardar el documento de Word con los datos reemplazados
    doc.save(output_docx_name)

    # Convertir el documento de Word a PDF usando docx2pdf
    convert(output_docx_name, output_pdf_name)

    # Leer el archivo PDF y convertirlo a base64
    with open(output_pdf_name, 'rb') as f:
        pdf_content = f.read()
        pdf_base64 = base64.b64encode(pdf_content).decode('utf-8')

    # Enviar la respuesta JSON con el archivo PDF en base64
    return JsonResponse({
        'response': 'Decomisos consultados correctamente',
        'Message': {
            'data': {
                'pdf_base64': pdf_base64
            }
        }
    }, status=200)
